from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# Создание Word-документа и запись текста
doc = Document()
doc.add_paragraph("Hello Python")
# Сохранение документа
doc.save("hello_python.docx")
# Чтение файла и выделение текста
loaded_doc = Document("hello_python.docx")
bold_text = ""
for paragraph in loaded_doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text
# Вывод на экран текста
print("Bold Text:", bold_text)
# Создание нового Word-документа с измененным шрифтом и размером
new_doc = Document()
paragraph = new_doc.add_paragraph("New paragraph with different font.")
run = paragraph.run
